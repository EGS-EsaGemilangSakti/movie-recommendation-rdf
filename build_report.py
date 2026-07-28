from pathlib import Path
from datetime import date

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips


ROOT = Path(__file__).resolve().parent
OUTPUT = ROOT / "Laporan_Dokumentasi_Movie_Recommendation_RDF.docx"
DIAGRAM = ROOT / ".report_architecture.png"

NAVY = "203748"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "263238"
GRAY = "5F6B73"
LIGHT = "F4F6F9"
PALE_BLUE = "E8EEF5"
WHITE = "FFFFFF"
GOLD = "A97800"
RED = "9B1C1C"
GREEN = "2F6B4F"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths_dxa, indent=120):
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[idx]
            cell.width = Twips(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_run_font(run, name="Calibri", size=None, color=INK, bold=None, italic=None):
    run.font.name = name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), name)
    r_fonts.set(qn("w:hAnsi"), name)
    r_fonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_field(paragraph, instruction, display=""):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = display
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    return run


def configure_numbering(doc):
    numbering = doc.part.numbering_part.element
    existing_abs = [
        int(x.get(qn("w:abstractNumId")))
        for x in numbering.findall(qn("w:abstractNum"))
    ]
    existing_num = [
        int(x.get(qn("w:numId")))
        for x in numbering.findall(qn("w:num"))
    ]
    next_abs = max(existing_abs or [0]) + 1
    next_num = max(existing_num or [0]) + 1

    def add_definition(abstract_id, num_id, fmt, text, font=None):
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        lvl.append(num_fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text)
        lvl.append(lvl_text)
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "tab")
        lvl.append(suff)
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "540")
        tabs.append(tab)
        p_pr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "540")
        ind.set(qn("w:hanging"), "280")
        p_pr.append(ind)
        lvl.append(p_pr)
        if font:
            r_pr = OxmlElement("w:rPr")
            fonts = OxmlElement("w:rFonts")
            fonts.set(qn("w:ascii"), font)
            fonts.set(qn("w:hAnsi"), font)
            r_pr.append(fonts)
            lvl.append(r_pr)
        abstract.append(lvl)
        numbering.append(abstract)

        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abstract_ref = OxmlElement("w:abstractNumId")
        abstract_ref.set(qn("w:val"), str(abstract_id))
        num.append(abstract_ref)
        numbering.append(num)

    add_definition(next_abs, next_num, "bullet", "•", "Arial")
    add_definition(next_abs + 1, next_num + 1, "decimal", "%1.")
    add_definition(next_abs + 2, next_num + 2, "decimal", "%1.")
    add_definition(next_abs + 3, next_num + 3, "decimal", "%1.")
    return next_num, (next_num + 1, next_num + 2, next_num + 3)


def apply_num(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_el])


def add_bullet(doc, text, bullet_num_id):
    p = doc.add_paragraph()
    apply_num(p, bullet_num_id)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.208
    set_run_font(p.add_run(text), size=11)
    return p


def add_number(doc, text, decimal_num_id):
    p = doc.add_paragraph()
    apply_num(p, decimal_num_id)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.208
    set_run_font(p.add_run(text), size=11)
    return p


def add_toc_entry(doc, title, page):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.15
    p_pr = p._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:leader"), "dot")
    tab.set(qn("w:pos"), "9000")
    tabs.append(tab)
    p_pr.append(tabs)
    set_run_font(p.add_run(f"{title}\t{page}"), size=10.5,
                 bold=title.startswith("BAB"))
    return p


def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.333
    if bold_lead and text.startswith(bold_lead):
        set_run_font(p.add_run(bold_lead), size=11, bold=True)
        set_run_font(p.add_run(text[len(bold_lead):]), size=11)
    else:
        set_run_font(p.add_run(text), size=11)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    set_run_font(run, size={1: 16, 2: 13, 3: 12}[level],
                 color={1: BLUE, 2: BLUE, 3: DARK_BLUE}[level], bold=True)
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.18)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.0
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F3F5F7")
    p_pr.append(shd)
    run = p.add_run(text)
    set_run_font(run, name="Consolas", size=9, color="263238")
    return p


def add_callout(doc, title, text, color=BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_run_font(r, size=10.5, color=color, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.15
    set_run_font(p2.add_run(text), size=10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc, headers, rows, widths, font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, PALE_BLUE)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(header), size=9.5, color=DARK_BLUE, bold=True)
    for row_data in rows:
        row = table.add_row()
        for idx, value in enumerate(row_data):
            cell = row.cells[idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            if idx == 0 and len(headers) > 2:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_run_font(p.add_run(str(value)), size=font_size)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def create_architecture_diagram(path):
    img = Image.new("RGB", (1500, 930), "white")
    draw = ImageDraw.Draw(img)
    font_path = "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"
    bold_path = "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"
    try:
        font = ImageFont.truetype(font_path, 30)
        small = ImageFont.truetype(font_path, 24)
        bold = ImageFont.truetype(bold_path, 31)
        title = ImageFont.truetype(bold_path, 38)
    except OSError:
        font = ImageFont.load_default(size=30)
        small = ImageFont.load_default(size=24)
        bold = ImageFont.load_default(size=31)
        title = ImageFont.load_default(size=38)
    draw.text((750, 42), "Arsitektur Movie Recommendation RDF", font=title,
              fill=(32, 55, 72), anchor="ma")

    boxes = [
        ((60, 175, 385, 355), "Dataset CSV", "Movies, Actors,\nDirectors, Countries", "E8EEF5"),
        ((590, 175, 910, 355), "rdf.py / RDFLib", "Transformasi data\nmenjadi triple RDF", "F4F6F9"),
        ((1115, 175, 1440, 355), "Turtle (.ttl)", "Knowledge graph\nhasil serialisasi", "E8EEF5"),
        ((1115, 555, 1440, 735), "Apache Fuseki", "Dataset /movies\nSPARQL endpoint", "FFF4D8"),
        ((590, 555, 910, 735), "app.py", "Streamlit +\nSPARQLWrapper", "E7F2EC"),
        ((60, 555, 385, 735), "Pengguna", "Pilih film, lihat detail,\ndan eksplorasi statistik", "F4F6F9"),
    ]
    for (x1, y1, x2, y2), head, body, fill in boxes:
        rgb = tuple(int(fill[i:i+2], 16) for i in (0, 2, 4))
        draw.rounded_rectangle((x1, y1, x2, y2), radius=24, fill=rgb,
                               outline=(46, 116, 181), width=4)
        draw.text(((x1+x2)//2, y1+36), head, font=bold, fill=(31, 77, 120), anchor="ma")
        draw.multiline_text(((x1+x2)//2, y1+92), body, font=small,
                            fill=(38, 50, 56), anchor="ma", align="center", spacing=8)

    def arrow(start, end, label):
        draw.line([start, end], fill=(46, 116, 181), width=7)
        ex, ey = end
        sx, sy = start
        if abs(ex-sx) > abs(ey-sy):
            direction = 1 if ex > sx else -1
            pts = [(ex, ey), (ex-18*direction, ey-12), (ex-18*direction, ey+12)]
        else:
            direction = 1 if ey > sy else -1
            pts = [(ex, ey), (ex-12, ey-18*direction), (ex+12, ey-18*direction)]
        draw.polygon(pts, fill=(46, 116, 181))
        mx, my = (sx+ex)//2, (sy+ey)//2
        draw.text((mx, my-18), label, font=font, fill=(95, 107, 115), anchor="ms")

    arrow((385, 265), (590, 265), "dibaca")
    arrow((910, 265), (1115, 265), "serialisasi")
    arrow((1277, 355), (1277, 555), "unggah")
    arrow((1115, 645), (910, 645), "JSON/SPARQL")
    arrow((590, 645), (385, 645), "antarmuka")
    draw.text((750, 840),
              "TMDB API melengkapi hasil dengan poster; movie_list.pkl berfungsi sebagai cache opsional.",
              font=small, fill=(95, 107, 115), anchor="ma")
    img.save(path)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_after = Pt(8)
normal.paragraph_format.line_spacing = 1.333
for style_name, size, color, before, after in [
    ("Heading 1", 16, BLUE, 18, 10),
    ("Heading 2", 13, BLUE, 12, 6),
    ("Heading 3", 12, DARK_BLUE, 8, 4),
]:
    style = styles[style_name]
    style.font.name = "Calibri"
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

if "Code Block" not in [s.name for s in styles]:
    code_style = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    code_style.font.name = "Consolas"
    code_style.font.size = Pt(9)

bullet_id, decimal_ids = configure_numbering(doc)

settings = doc.settings._element
update_fields = OxmlElement("w:updateFields")
update_fields.set(qn("w:val"), "true")
settings.append(update_fields)

# Running header and footer
header = section.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_run_font(hp.add_run("Dokumentasi Teknis | Movie Recommendation RDF"),
             size=8.5, color=GRAY, italic=True)
footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run_font(fp.add_run("Halaman "), size=9, color=GRAY)
add_field(fp, "PAGE", "1")

# Cover
for _ in range(4):
    doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(18)
set_run_font(p.add_run("LAPORAN DOKUMENTASI PROYEK"), size=13, color=GOLD, bold=True)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(10)
set_run_font(p.add_run("MOVIE RECOMMENDATION SYSTEM"), size=28, color=NAVY, bold=True)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(8)
set_run_font(p.add_run("Menggunakan RDF dan SPARQL"), size=20, color=DARK_BLUE, bold=True)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(52)
set_run_font(p.add_run("Analisis Repository, Arsitektur, Implementasi, dan Panduan Operasional"),
             size=12.5, color=GRAY, italic=True)
add_callout(doc, "Ruang lingkup laporan",
            "Dokumen ini disusun berdasarkan working tree repository per 29 Juli 2026, "
            "termasuk pembaruan kompatibilitas Python 3.14, perbaikan kueri SPARQL, "
            "integrasi Fuseki, dan dokumentasi operasional.", color=DARK_BLUE)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(36)
set_run_font(p.add_run("Disusun untuk dokumentasi akademik dan teknis"), size=11, color=GRAY)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run_font(p.add_run("Team Itihaad"), size=12, color=NAVY, bold=True)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run_font(p.add_run("Juli 2026"), size=11, color=GRAY)
doc.add_page_break()

# Kata Pengantar
add_heading(doc, "KATA PENGANTAR", 1)
add_body(doc, "Puji syukur ke hadirat Tuhan Yang Maha Esa karena laporan dokumentasi proyek "
         "Movie Recommendation System Using RDF and SPARQL ini dapat disusun. Laporan ini "
         "bertujuan memberikan gambaran menyeluruh mengenai latar belakang, rancangan, "
         "struktur repository, implementasi, pengujian, serta tata cara menjalankan sistem.")
add_body(doc, "Penyusunan laporan dilakukan melalui penelaahan langsung terhadap source code, "
         "notebook, dataset CSV, artefak RDF/Turtle, konfigurasi dependensi, dan hasil kueri "
         "pada Apache Jena Fuseki. Dengan pendekatan tersebut, uraian dalam laporan merefleksikan "
         "kondisi repository aktual dan tidak hanya mengulang informasi pada README.")
add_body(doc, "Penulis menyadari bahwa proyek masih dapat dikembangkan, terutama pada aspek "
         "personalisasi rekomendasi, pengujian otomatis, pengelolaan konfigurasi, dan deployment. "
         "Oleh sebab itu, laporan turut memuat evaluasi kritis dan rekomendasi pengembangan "
         "lanjutan. Semoga dokumen ini bermanfaat bagi pengembang, penguji, dosen, mahasiswa, "
         "dan pihak lain yang mempelajari penerapan Semantic Web pada domain film.")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p.paragraph_format.space_before = Pt(24)
set_run_font(p.add_run("Jakarta, 29 Juli 2026"), size=11)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_run_font(p.add_run("Team Itihaad"), size=11, bold=True)
doc.add_page_break()

# TOC
add_heading(doc, "DAFTAR ISI", 1)
for toc_title, toc_page in [
    ("KATA PENGANTAR", 2),
    ("DAFTAR ISI", 3),
    ("RINGKASAN EKSEKUTIF", 4),
    ("BAB I PENDAHULUAN", 5),
    ("BAB II LANDASAN TEORI", 7),
    ("BAB III ANALISIS REPOSITORY", 9),
    ("BAB IV ANALISIS DAN PERANCANGAN SISTEM", 12),
    ("BAB V IMPLEMENTASI SISTEM", 15),
    ("BAB VI PENGUJIAN DAN HASIL", 18),
    ("BAB VII EVALUASI PROYEK", 20),
    ("BAB VIII PANDUAN OPERASIONAL", 22),
    ("BAB IX PENUTUP", 24),
    ("DAFTAR PUSTAKA", 25),
    ("LAMPIRAN A CONTOH KUERI SPARQL", 26),
    ("LAMPIRAN B KAMUS DATA FILM", 27),
]:
    add_toc_entry(doc, toc_title, toc_page)
doc.add_page_break()

# Executive summary
add_heading(doc, "RINGKASAN EKSEKUTIF", 1)
add_body(doc, "Movie Recommendation System Using RDF and SPARQL merupakan aplikasi eksplorasi "
         "film berbasis knowledge graph. Sistem mengubah data tabular menjadi RDF, memublikasikan "
         "graph melalui Apache Jena Fuseki, lalu menjalankan kueri SPARQL dari aplikasi Streamlit. "
         "Poster film diperoleh secara opsional dari TMDB API.")
add_callout(doc, "Temuan utama",
            "Secara fungsional proyek telah mendukung pencarian detail film, rekomendasi berbasis "
            "aktor atau sutradara yang sama, serta eksplorasi statistik. Namun, istilah "
            "\"rekomendasi\" masih bersifat content/relationship-based sederhana dan belum "
            "mencakup personalisasi pengguna atau evaluasi model machine learning.", color=BLUE)
add_table(doc,
          ["Aspek", "Ringkasan"],
          [
              ("Data", "180 baris film, 179 judul unik, 389 aktor, 10 sutradara, dan 19 negara."),
              ("Knowledge graph", "Empat kelas utama: Movie, Actor, Director, dan Country."),
              ("Layanan data", "Apache Jena Fuseki pada endpoint /movies/sparql."),
              ("Antarmuka", "Streamlit dengan pencarian film, detail, rekomendasi relasional, dan enam kueri agregat."),
              ("Kesiapan", "Layak untuk demonstrasi akademik; memerlukan penguatan untuk produksi."),
          ],
          [2100, 7260])

# BAB I
doc.add_page_break()
add_heading(doc, "BAB I PENDAHULUAN", 1)
add_heading(doc, "1.1 Latar Belakang", 2)
add_body(doc, "Data film lazim tersedia dalam bentuk tabel, JSON, atau basis data relasional. "
         "Format tersebut efektif untuk penyimpanan, tetapi hubungan antarpelaku, sutradara, "
         "negara, genre, dan film sering kali tersebar di berbagai kolom atau tabel. Pendekatan "
         "Semantic Web menawarkan model graph yang menempatkan hubungan sebagai bagian inti data.")
add_body(doc, "RDF merepresentasikan pengetahuan sebagai triple subjek-predikat-objek, sedangkan "
         "SPARQL menyediakan bahasa kueri untuk menelusuri pola hubungan pada graph. Dalam proyek "
         "ini, domain film digunakan sebagai studi kasus untuk menunjukkan bagaimana data CSV "
         "dikonversi menjadi knowledge graph, disajikan melalui endpoint HTTP, dan diakses oleh "
         "aplikasi web interaktif.")
add_heading(doc, "1.2 Rumusan Masalah", 2)
for text in [
    "Bagaimana mengubah data film yang semula tabular menjadi representasi RDF yang terstruktur?",
    "Bagaimana memodelkan hubungan film dengan aktor, sutradara, genre, dan negara produksi?",
    "Bagaimana menjalankan pencarian dan agregasi terhadap knowledge graph menggunakan SPARQL?",
    "Bagaimana menyajikan hasil kueri melalui antarmuka web yang mudah digunakan?",
    "Apa kelebihan, keterbatasan, dan arah pengembangan repository pada kondisi saat ini?",
]:
    add_bullet(doc, text, bullet_id)
add_heading(doc, "1.3 Tujuan", 2)
for text in [
    "Membangun representasi knowledge graph untuk domain film.",
    "Mendemonstrasikan pemanfaatan RDFLib untuk membentuk dan menyerialisasi RDF.",
    "Menyediakan endpoint SPARQL menggunakan Apache Jena Fuseki.",
    "Membangun antarmuka Streamlit untuk pencarian, eksplorasi, dan rekomendasi relasional.",
    "Mendokumentasikan struktur repository, modul, alur kerja, hasil uji, dan peluang pengembangan.",
]:
    add_bullet(doc, text, bullet_id)
add_heading(doc, "1.4 Ruang Lingkup", 2)
add_body(doc, "Laporan membahas working tree repository lokal per 29 Juli 2026. Analisis mencakup "
         "file Python, notebook, dataset, file Turtle, dependensi, konfigurasi Fuseki, kueri "
         "SPARQL, serta integrasi TMDB. Laporan tidak menilai akurasi rekomendasi personal karena "
         "repository belum memiliki profil pengguna, data rating pengguna, algoritma collaborative "
         "filtering, atau metrik evaluasi rekomendasi.")
add_heading(doc, "1.5 Metode Penyusunan", 2)
for text in [
    "Inventarisasi struktur file dan riwayat repository.",
    "Pembacaan statis terhadap app.py, rdf.py, generate_movie_list.py, dan varian aplikasi.",
    "Analisis notebook preprocessing dan notebook konstruksi RDF.",
    "Perhitungan ringkas karakteristik dataset dan entitas RDF.",
    "Validasi kueri terhadap endpoint Fuseki lokal.",
    "Perbandingan implementasi dengan dokumentasi resmi teknologi terkait.",
]:
    add_number(doc, text, decimal_ids[0])

# BAB II
doc.add_page_break()
add_heading(doc, "BAB II LANDASAN TEORI", 1)
add_heading(doc, "2.1 Semantic Web dan Knowledge Graph", 2)
add_body(doc, "Semantic Web bertujuan membuat data dapat dipahami dan diproses secara lebih "
         "terstruktur oleh mesin. Knowledge graph mengorganisasi entitas dan relasi dalam bentuk "
         "graph sehingga pertanyaan yang berorientasi hubungan dapat dijawab tanpa bergantung "
         "pada skema tabel yang kaku.")
add_heading(doc, "2.2 Resource Description Framework (RDF)", 2)
add_body(doc, "RDF adalah model data graph yang tersusun atas triple: subjek sebagai sumber daya, "
         "predikat sebagai jenis hubungan atau atribut, dan objek sebagai sumber daya lain atau "
         "literal. W3C RDF 1.1 juga mendefinisikan penggunaan IRI, literal bertipe, serta konsep "
         "dataset yang dapat memuat lebih dari satu graph [1].")
add_code(doc, 'ex:Movie10054  ex:title       "Spy Kids" .\n'
              'ex:Movie10054  ex:directedBy  ex:DirectorRobert%20Rodriguez .')
add_heading(doc, "2.3 RDFS dan XSD", 2)
add_body(doc, "RDF Schema (RDFS) dipakai untuk menyatakan kelas, properti, domain, dan range. "
         "XML Schema Datatypes (XSD) memberi tipe literal seperti xsd:string, xsd:date, "
         "xsd:integer, dan xsd:float. Repository menggunakan keduanya agar makna atribut dan "
         "tipe nilai lebih eksplisit.")
add_heading(doc, "2.4 SPARQL", 2)
add_body(doc, "SPARQL adalah bahasa kueri untuk RDF. SPARQL 1.1 mendukung pencocokan pola triple, "
         "filter, agregasi, pengelompokan, pengurutan, subkueri, dan berbagai fungsi operator [2]. "
         "Repository menggunakan SELECT, COUNT, GROUP BY, ORDER BY, FILTER, OPTIONAL, BIND, dan "
         "LIMIT untuk membentuk fitur aplikasi.")
add_heading(doc, "2.5 Apache Jena Fuseki", 2)
add_body(doc, "Fuseki merupakan server SPARQL yang memublikasikan data RDF melalui protokol HTTP. "
         "Fuseki dapat dijalankan sebagai server mandiri dan terintegrasi dengan TDB untuk "
         "penyimpanan persisten [3]. Pada proyek ini, dataset diberi nama movies sehingga endpoint "
         "kuerinya adalah http://localhost:3030/movies/sparql.")
add_heading(doc, "2.6 Sistem Rekomendasi", 2)
add_body(doc, "Sistem rekomendasi dapat menggunakan collaborative filtering, content-based "
         "filtering, pendekatan hybrid, atau hubungan semantik. Repository ini merekomendasikan "
         "film lain yang memiliki aktor atau sutradara sama. Oleh karena itu, mekanismenya paling "
         "tepat disebut rekomendasi relasional berbasis konten, bukan rekomendasi personal.")

# BAB III
doc.add_page_break()
add_heading(doc, "BAB III ANALISIS REPOSITORY", 1)
add_heading(doc, "3.1 Identitas dan Kondisi Repository", 2)
add_table(doc, ["Atribut", "Nilai"], [
    ("Nama proyek", "Movie Recommendation System Using RDF and SPARQL"),
    ("Bahasa utama", "Python"),
    ("Runtime target", "Python 3.14"),
    ("Branch aktif", "main"),
    ("Versi terdokumentasi", "v1.0.0 pada riwayat Git; working tree memuat pembaruan sesudahnya"),
    ("Aplikasi utama", "app.py"),
    ("Endpoint default", "http://localhost:3030/movies/sparql"),
    ("Tanggal analisis", "29 Juli 2026"),
], [2100, 7260])
add_callout(doc, "Catatan working tree",
            "README.md, app.py, dan rdf.py sedang memuat perubahan lokal yang belum tercermin "
            "pada commit HEAD. Laporan ini mendokumentasikan keadaan file aktual, bukan hanya "
            "snapshot commit terakhir.", color=GOLD)
add_heading(doc, "3.2 Struktur Direktori", 2)
add_code(doc,
""".
├── app.py
├── app22.py
├── o_app.py
├── rdf.py
├── generate_movie_list.py
├── requirements.txt
├── .python-version
├── output_with_classes_and_properties.ttl
├── Preproccessing.ipynb
├── Project Code.ipynb
├── Report.pdf
└── Datasets/
    ├── Movies_less.csv
    ├── Actors.csv
    ├── Directors.csv
    └── Countries.csv""")
add_heading(doc, "3.3 Penjelasan File dan Modul", 2)
add_table(doc, ["No.", "File/Modul", "Peran"], [
    ("1", "app.py", "Aplikasi Streamlit utama; menjalankan kueri SPARQL, menyajikan detail dan rekomendasi, serta mengambil poster TMDB."),
    ("2", "rdf.py", "Membaca CSV, mendefinisikan kelas/properti, membentuk triple RDF, dan menyimpan graph dalam Turtle."),
    ("3", "generate_movie_list.py", "Membuat cache movie_list.pkl dari kolom title agar dropdown dapat dimuat cepat."),
    ("4", "Preproccessing.ipynb", "Notebook pembersihan dan transformasi dataset TMDB; 39 sel kode dan 12 sel Markdown."),
    ("5", "Project Code.ipynb", "Notebook ringkas yang mendokumentasikan konstruksi RDF; satu sel kode utama."),
    ("6", "output_with_classes_and_properties.ttl", "Artefak knowledge graph berformat Turtle dengan ukuran sekitar 256 KiB."),
    ("7", "Datasets/*.csv", "Sumber data film, aktor, sutradara, dan negara."),
    ("8", "app22.py / o_app.py", "Varian implementasi lama dengan namespace dan pola kueri berbeda; bukan entry point utama."),
    ("9", "requirements.txt", "Pin versi dependensi Python yang ditargetkan kompatibel dengan Python 3.14."),
    ("10", "README.md", "Panduan instalasi, Fuseki, Java 21, konfigurasi, eksekusi, dan troubleshooting."),
], [650, 2150, 6560], font_size=8.9)
add_heading(doc, "3.4 Karakteristik Dataset", 2)
add_table(doc, ["Dataset", "Jumlah", "Keterangan"], [
    ("Movies_less.csv", "180 baris / 179 judul unik", "Film berbahasa Inggris, rentang rilis 1966-2016."),
    ("Actors.csv", "389 aktor", "Nama, tanggal lahir, negara, dan estimasi kekayaan."),
    ("Directors.csv", "10 sutradara", "Nama, tanggal lahir, negara, dan estimasi kekayaan."),
    ("Countries.csv", "19 negara", "Benua, populasi, ibu kota, dan tanggal kemerdekaan."),
    ("Genre", "18 kategori unik", "Action hingga Western."),
    ("Turtle", "5.018 triple pada Fuseki", "Hasil hitung endpoint lokal sebelum regenerasi terbaru."),
], [2350, 1900, 5110])
add_body(doc, "Terdapat satu judul duplikat atau baris yang berkonvergensi menjadi 179 entitas film "
         "unik pada artefak RDF. Seluruh film pada dataset ringkas menggunakan nilai bahasa "
         "English. Genre yang ditemukan meliputi Action, Adventure, Animation, Comedy, Crime, "
         "Documentary, Drama, Family, Fantasy, History, Horror, Music, Mystery, Romance, Science "
         "Fiction, Thriller, War, dan Western.")
add_heading(doc, "3.5 Dependensi", 2)
add_table(doc, ["Paket", "Versi", "Fungsi"], [
    ("Streamlit", "1.60.0", "Antarmuka web interaktif."),
    ("RDFLib", "7.6.0", "Pembuatan graph, URI, literal, namespace, dan serialisasi Turtle."),
    ("SPARQLWrapper", "2.0.0", "Klien Python untuk endpoint SPARQL."),
    ("pandas", "3.0.3", "Pembacaan dan transformasi data CSV."),
    ("isodate", "0.7.2", "Validasi tanggal ISO 8601."),
    ("requests", "2.34.2", "Permintaan HTTP ke TMDB API."),
], [2300, 1300, 5760])

# BAB IV
doc.add_page_break()
add_heading(doc, "BAB IV ANALISIS DAN PERANCANGAN SISTEM", 1)
add_heading(doc, "4.1 Kebutuhan Fungsional", 2)
for text in [
    "Sistem menampilkan daftar judul film untuk dipilih pengguna.",
    "Sistem mengambil detail film dari knowledge graph melalui SPARQL.",
    "Sistem menampilkan rekomendasi film dengan sutradara atau aktor yang sama.",
    "Sistem menyediakan enam pertanyaan agregat untuk budget, runtime, revenue, genre, dan negara.",
    "Sistem menampilkan poster melalui TMDB API ketika API key tersedia.",
    "Sistem memberi pesan yang dapat ditindaklanjuti jika Fuseki tidak dapat diakses.",
]:
    add_bullet(doc, text, bullet_id)
add_heading(doc, "4.2 Kebutuhan Nonfungsional", 2)
for text in [
    "Kompatibilitas dengan Python 3.14 dan Java 21+ untuk Jena 6.",
    "Konfigurasi endpoint dan API key melalui environment variable.",
    "Waktu tunggu jaringan terbatas untuk mencegah aplikasi menggantung.",
    "Path data relatif terhadap lokasi source code agar tidak bergantung pada working directory.",
    "Data RDF dapat disimpan persisten melalui TDB2 pada Fuseki.",
]:
    add_bullet(doc, text, bullet_id)
add_heading(doc, "4.3 Arsitektur Sistem", 2)
create_architecture_diagram(DIAGRAM)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run()
run.add_picture(str(DIAGRAM), width=Inches(6.3))
p.paragraph_format.space_after = Pt(3)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run_font(p.add_run("Gambar 1. Arsitektur logis sistem"), size=9.5, color=GRAY, italic=True)
add_body(doc, "Arsitektur memisahkan pipeline pembentukan data dari aplikasi konsumsi. rdf.py "
         "menghasilkan file Turtle yang kemudian diunggah ke Fuseki. app.py tidak membaca Turtle "
         "secara langsung, tetapi mengirim kueri ke endpoint SPARQL. Pemisahan ini memungkinkan "
         "graph diakses oleh lebih dari satu klien dan dikelola secara persisten.")
add_heading(doc, "4.4 Model Ontologi", 2)
add_table(doc, ["Kelas", "Representasi", "Relasi/Atribut Penting"], [
    ("Movie", "ex:Movie<ID>", "movie_id, title, overview, runtime, budget, revenue, release_date, genre, original_language."),
    ("Actor", "ex:Actor<Nama>", "title, date_of_birth, country, net_worth; terhubung melalui hasActor."),
    ("Director", "ex:Director<Nama>", "title, date_of_birth, country, net_worth; terhubung melalui directedBy."),
    ("Country", "ex:Country<Nama>", "title, continent, population, capital, independence_date; terhubung melalui producedIn."),
], [1500, 2350, 5510], font_size=8.9)
add_heading(doc, "4.5 Alur Kerja", 2)
for text in [
    "Data mentah diproses pada notebook preprocessing menjadi CSV terstruktur.",
    "rdf.py membaca keempat CSV dan membentuk URI serta literal bertipe.",
    "Graph diserialisasi sebagai output_with_classes_and_properties.ttl.",
    "File Turtle diunggah ke dataset movies pada Apache Jena Fuseki.",
    "Streamlit membaca daftar judul dari cache atau CSV.",
    "Pilihan pengguna diubah menjadi literal SPARQL yang telah di-escape.",
    "Fuseki menjalankan kueri dan mengembalikan binding JSON.",
    "Aplikasi memformat detail, statistik, rekomendasi, dan poster.",
]:
    add_number(doc, text, decimal_ids[1])

# BAB V
doc.add_page_break()
add_heading(doc, "BAB V IMPLEMENTASI SISTEM", 1)
add_heading(doc, "5.1 Tahap Preprocessing", 2)
add_body(doc, "Preproccessing.ipynb membaca dataset TMDB 5000 Movies dan Credits, menggabungkan "
         "keduanya berdasarkan title, menghapus nilai kosong dan duplikat, serta mengurai kolom "
         "JSON seperti genres, production_countries, cast, dan crew. Hasilnya diratakan menjadi "
         "kolom genre_1 sampai genre_5, Actor_1 sampai Actor_3, Country_1 sampai Country_5, dan "
         "Director. Notebook menggunakan kernel Python 3.14 pada metadata terbaru.")
add_heading(doc, "5.2 Pembentukan RDF oleh rdf.py", 2)
add_body(doc, "Modul rdf.py membuat Graph RDFLib, mengikat namespace ex, RDFS, dan XSD, lalu "
         "mendefinisikan empat kelas. Fungsi create_uri membersihkan dan melakukan URL encoding "
         "pada nilai sebelum membentuk URIRef. Fungsi is_valid_date memanfaatkan isodate untuk "
         "menentukan apakah tanggal dapat ditulis sebagai xsd:date.")
add_code(doc,
"""movie_uri = create_uri(ex.Movie, str(row["movie_id"]))
g.add((movie_uri, RDF.type, ex.Movie))
g.add((movie_uri, ex.movie_id,
       Literal(row["movie_id"], datatype=XSD.integer)))""")
add_body(doc, "Untuk setiap film, modul menambahkan literal deskriptif dan relasi ke sutradara, "
         "aktor, genre, serta negara. Entitas aktor, sutradara, dan negara kemudian diperkaya "
         "dengan atribut masing-masing. Graph disimpan dengan encoding UTF-8 dalam format Turtle.")
add_heading(doc, "5.3 Cache Judul Film", 2)
add_body(doc, "generate_movie_list.py memilih kolom title, menghapus nilai kosong dan duplikat, "
         "mereset indeks, lalu menyimpan DataFrame sebagai movie_list.pkl. app.py memperlakukan "
         "cache ini sebagai opsional: jika file tidak tersedia, aplikasi membaca kolom title "
         "langsung dari Movies_less.csv.")
add_heading(doc, "5.4 Integrasi Apache Jena Fuseki", 2)
add_body(doc, "Fuseki 6.1.0 memerlukan Java 21 atau lebih baru. Distribusi yang digunakan harus "
         "bernama apache-jena-fuseki-6.1.0.zip, bukan source release. Dataset dibuat dengan nama "
         "movies dan disarankan menggunakan TDB2 agar data persisten. Endpoint default aplikasi "
         "dapat diubah melalui environment variable SPARQL_ENDPOINT.")
add_code(doc, '$env:SPARQL_ENDPOINT = "http://localhost:3030/movies/sparql"\n'
              '.\\fuseki-server.bat')
add_heading(doc, "5.5 Modul Aplikasi app.py", 2)
add_table(doc, ["Fungsi/Bagian", "Penjelasan"], [
    ("fetch_poster(movie_id)", "Mengambil poster TMDB melalui HTTPS, memakai timeout 10 detik, cache Streamlit, dan placeholder saat gagal."),
    ("run_sparql_query(query)", "Membuat klien SPARQLWrapper, meminta JSON, menetapkan timeout 15 detik, dan menampilkan pesan jika Fuseki gagal."),
    ("sparql_string(value)", "Membentuk literal xsd:string yang aman menggunakan RDFLib untuk mencegah kerusakan sintaks kueri."),
    ("binding_value(...)", "Mengambil value dari binding JSON dengan nilai default yang konsisten."),
    ("Pemuatan movie list", "Membaca movie_list.pkl atau fallback CSV; menghapus nilai kosong/duplikat dan mengurutkan judul."),
    ("Rendering hasil", "Menggunakan komponen Streamlit columns, image, write, caption, radio, selectbox, dan button."),
], [2700, 6660], font_size=8.9)
add_heading(doc, "5.6 Fitur Pencarian Film", 2)
add_body(doc, "Pengguna memilih judul dari dropdown. Kueri detail mengambil overview, tanggal "
         "rilis, runtime, budget, revenue, sutradara, aktor, negara, bahasa, genre, dan movie ID. "
         "Nama sutradara dan negara diperoleh dari entitas terkait. OPTIONAL dan BIND digunakan "
         "untuk menjaga kompatibilitas dengan file Turtle lama yang menyimpan ID hanya dalam URI.")
add_heading(doc, "5.7 Fitur Rekomendasi Relasional", 2)
add_body(doc, "Rekomendasi sutradara mencari hingga lima film lain yang memiliki directedBy sama. "
         "Rekomendasi aktor mencari hingga lima film lain yang berbagi minimal satu aktor. FILTER "
         "mencegah film yang sedang dipilih muncul kembali. Mekanisme ini transparan dan mudah "
         "dijelaskan, tetapi belum menghitung skor kemiripan atau preferensi pengguna.")
add_heading(doc, "5.8 Pertanyaan SPARQL Siap Pakai", 2)
add_table(doc, ["No.", "Pertanyaan", "Operasi Utama"], [
    ("1", "Top 10 Movies by Budget", "ORDER BY DESC(?budget)"),
    ("2", "Top 10 Movies by Runtime", "ORDER BY DESC(?runtime)"),
    ("3", "Top 10 Highest Revenue Movies", "ORDER BY DESC(?revenue)"),
    ("4", "Top 10 Most Popular Genres", "COUNT, GROUP BY genre"),
    ("5", "Top 10 Most Common Countries", "COUNT, GROUP BY country"),
    ("6", "Movies with the Most Countries of Origin", "COUNT, GROUP BY movie"),
], [650, 4700, 4010])
add_heading(doc, "5.9 Integrasi TMDB", 2)
add_body(doc, "TMDB API dipakai hanya untuk poster. API key tidak lagi ditulis pada source code, "
         "melainkan dibaca dari TMDB_API_KEY. Tanpa key, fitur inti tetap berjalan menggunakan "
         "placeholder. Pendekatan ini mengurangi risiko kebocoran kredensial dan menjaga aplikasi "
         "tetap usable pada lingkungan offline sebagian.")

# BAB VI
doc.add_page_break()
add_heading(doc, "BAB VI PENGUJIAN DAN HASIL", 1)
add_heading(doc, "6.1 Strategi Pengujian", 2)
add_body(doc, "Pengujian yang dilakukan untuk laporan ini bersifat verifikasi teknis terhadap "
         "working tree. Cakupannya meliputi validitas sintaks Python, konsistensi skema RDF dengan "
         "kueri aplikasi, hasil kueri endpoint, validitas JSON notebook, dan pemeriksaan diff Git.")
add_heading(doc, "6.2 Hasil Uji Endpoint Fuseki", 2)
add_table(doc, ["Skenario", "Hasil", "Status"], [
    ("Hitung seluruh triple", "5.018 triple", "Lulus"),
    ("Detail film Spy Kids", "12 binding kombinasi aktor/genre", "Lulus"),
    ("Rekomendasi sutradara", "5 film", "Lulus"),
    ("Rekomendasi aktor", "4 film", "Lulus"),
    ("Top budget", "10 hasil", "Lulus"),
    ("Top negara", "10 hasil", "Lulus"),
    ("Jumlah negara per film", "10 hasil", "Lulus"),
], [3900, 3300, 2160])
add_heading(doc, "6.3 Hasil Uji Sintaks dan Struktur", 2)
for text in [
    "app.py, rdf.py, dan generate_movie_list.py berhasil dikompilasi dengan compileall pada container Python tanpa error sintaks.",
    "Preproccessing.ipynb dan Project Code.ipynb valid sebagai JSON dan menggunakan metadata kernel Python 3.14.",
    "git diff --check tidak menemukan whitespace error.",
    "Tidak ditemukan referensi Fivetran atau webhook pada source code repository.",
]:
    add_bullet(doc, text, bullet_id)
add_heading(doc, "6.4 Analisis Konsistensi Artefak", 2)
add_callout(doc, "Perhatian",
            "rdf.py terbaru telah menambahkan properti ex:movie_id, tetapi file Turtle di repository "
            "dibuat sebelum perubahan tersebut. app.py menyediakan fallback dengan mengekstrak ID "
            "dari URI ex:Movie<ID>. Untuk konsistensi penuh, jalankan ulang rdf.py dan unggah ulang "
            "Turtle ke Fuseki.", color=GOLD)
add_heading(doc, "6.5 Interpretasi Hasil", 2)
add_body(doc, "Hasil uji menunjukkan bahwa arsitektur utama bekerja: data tersedia di Fuseki, "
         "kueri detail dan agregat menghasilkan binding, serta aplikasi telah memiliki fallback "
         "untuk artefak RDF lama. Pengujian belum mencakup benchmark beban, uji antarmuka otomatis, "
         "keamanan endpoint, atau evaluasi kualitas rekomendasi.")

# BAB VII
doc.add_page_break()
add_heading(doc, "BAB VII EVALUASI PROYEK", 1)
add_heading(doc, "7.1 Kelebihan", 2)
for text in [
    "Pemisahan pipeline data, server SPARQL, dan antarmuka cukup jelas.",
    "Model graph tepat untuk merepresentasikan hubungan film-aktor-sutradara-negara.",
    "Kueri SPARQL mudah dibaca dan dapat dijelaskan dalam konteks akademik.",
    "Konfigurasi sensitif telah dipindahkan ke environment variable.",
    "Aplikasi tetap berfungsi ketika cache judul atau API poster tidak tersedia.",
    "Dokumentasi instalasi telah mencakup Java 21, Fuseki, dataset, dan troubleshooting.",
]:
    add_bullet(doc, text, bullet_id)
add_heading(doc, "7.2 Keterbatasan", 2)
for text in [
    "Rekomendasi belum personal dan tidak memiliki skor relevansi.",
    "Tidak tersedia unit test, integration test, atau UI test dalam repository.",
    "app22.py dan o_app.py mempertahankan namespace lama sehingga berpotensi membingungkan.",
    "Pipeline preprocessing bergantung pada dataset mentah TMDB yang tidak disertakan.",
    "File pickle tetap memiliki risiko keamanan jika berasal dari sumber tidak dipercaya.",
    "Fuseki harus dikonfigurasi manual dan belum tersedia konfigurasi deployment otomatis.",
    "Pemodelan country pada aktor/sutradara menggunakan literal, sedangkan range dideklarasikan ex:Country.",
    "vote_average dan vote_count didefinisikan sebagai properti tetapi belum ditulis ke graph pada rdf.py.",
]:
    add_bullet(doc, text, bullet_id)
add_heading(doc, "7.3 Risiko Teknis", 2)
add_table(doc, ["Risiko", "Dampak", "Mitigasi"], [
    ("Artefak RDF tidak sinkron", "Kueri dapat kehilangan field baru.", "Regenerasi Turtle dan re-upload terkontrol."),
    ("Fuseki tidak aktif", "Aplikasi tidak dapat mengambil data.", "Health check dan pesan error operasional."),
    ("API key bocor", "Penyalahgunaan layanan TMDB.", "Environment variable/secrets dan rotasi key."),
    ("Pickle tidak tepercaya", "Eksekusi kode saat deserialisasi.", "Gunakan cache lokal atau ganti format Parquet/CSV."),
    ("Varian kode lama", "Pemelihara menjalankan entry point salah.", "Arsipkan atau pindahkan ke folder legacy."),
], [2700, 3000, 3660], font_size=8.8)
add_heading(doc, "7.4 Rekomendasi Pengembangan", 2)
for text in [
    "Tambahkan model skor rekomendasi berbasis kemiripan genre, aktor, sutradara, dan atribut numerik.",
    "Tambahkan profil, rating, atau riwayat pengguna untuk personalisasi.",
    "Bangun test suite untuk generator RDF, kueri SPARQL, dan fungsi parsing binding.",
    "Sediakan Docker Compose untuk Fuseki dan aplikasi agar setup dapat direproduksi.",
    "Gunakan configuration file atau environment-specific settings untuk endpoint.",
    "Tambahkan CI yang menguji Python 3.14, validitas Turtle, dan kueri smoke test.",
    "Regenerasi Turtle setelah perubahan skema serta catat versi artefak.",
    "Pisahkan lapisan data, service, dan UI agar app.py lebih modular.",
]:
    add_number(doc, text, decimal_ids[2])

# BAB VIII
doc.add_page_break()
add_heading(doc, "BAB VIII PANDUAN OPERASIONAL", 1)
add_heading(doc, "8.1 Prasyarat", 2)
for text in [
    "Python 3.14 dan pip.",
    "Java 21 atau lebih baru.",
    "Distribusi biner apache-jena-fuseki-6.1.0.zip.",
    "Koneksi internet jika poster TMDB diperlukan.",
]:
    add_bullet(doc, text, bullet_id)
add_heading(doc, "8.2 Instalasi Python", 2)
add_code(doc, "py -3.14 -m venv .venv\n"
              ".\\.venv\\Scripts\\Activate.ps1\n"
              "python -m pip install --upgrade pip\n"
              "python -m pip install -r requirements.txt")
add_heading(doc, "8.3 Menjalankan Fuseki", 2)
add_code(doc, "winget install EclipseAdoptium.Temurin.21.JDK\n"
              "cd C:\\apache-jena-fuseki-6.1.0\n"
              ".\\fuseki-server.bat")
add_body(doc, "Buka http://localhost:3030, buat dataset movies dengan TDB2, kemudian unggah "
         "output_with_classes_and_properties.ttl ke default graph. Uji dengan COUNT(*) dan "
         "pastikan nilai total lebih besar dari nol.")
add_heading(doc, "8.4 Menjalankan Aplikasi", 2)
add_code(doc, '$env:TMDB_API_KEY = "API_KEY_TMDB_MILIK_ANDA"\n'
              'streamlit run app.py')
add_heading(doc, "8.5 Regenerasi Data", 2)
add_code(doc, "python generate_movie_list.py\npython rdf.py")
add_body(doc, "Setelah rdf.py dijalankan, bersihkan atau buat ulang dataset movies lalu unggah "
         "Turtle terbaru. Jangan mengunggah file yang sama berulang kali tanpa pengelolaan data.")
add_heading(doc, "8.6 Troubleshooting Ringkas", 2)
add_table(doc, ["Gejala", "Penyebab", "Tindakan"], [
    ("'java' is not recognized", "JDK belum terpasang/PATH belum terbaca.", "Pasang Temurin 21, buka terminal baru, lalu cek java -version."),
    ("fuseki-server.jar tidak ada", "Source release atau paket salah.", "Unduh apache-jena-fuseki-6.1.0.zip."),
    ("SPARQL gagal", "Fuseki mati, dataset salah, atau endpoint berbeda.", "Aktifkan Fuseki dan cek /movies/sparql."),
    ("Poster tidak tampil", "TMDB_API_KEY kosong/tidak valid.", "Set key pada terminal yang sama atau gunakan placeholder."),
    ("movie_list.pkl tidak ada", "Cache belum dibuat.", "Aplikasi memakai CSV; opsional jalankan generate_movie_list.py."),
], [2550, 3300, 3510], font_size=8.7)

# BAB IX
doc.add_page_break()
add_heading(doc, "BAB IX PENUTUP", 1)
add_heading(doc, "9.1 Kesimpulan", 2)
add_body(doc, "Repository berhasil mendemonstrasikan integrasi data film, RDF, SPARQL, Apache "
         "Jena Fuseki, dan Streamlit dalam satu alur yang dapat dijalankan. Empat kelas utama dan "
         "relasi semantik memungkinkan pencarian detail serta rekomendasi berbasis kesamaan aktor "
         "atau sutradara. Pengujian endpoint menunjukkan graph aktif dan kueri utama menghasilkan "
         "data sesuai rancangan.")
add_body(doc, "Dari perspektif akademik, proyek efektif sebagai demonstrasi Semantic Web dan "
         "knowledge graph. Dari perspektif produk, proyek masih merupakan aplikasi eksplorasi "
         "film dengan rekomendasi relasional sederhana. Kesiapan produksi memerlukan pengujian "
         "otomatis, deployment reproducible, konsistensi artefak, modularisasi, dan mekanisme "
         "rekomendasi yang dapat dievaluasi.")
add_heading(doc, "9.2 Saran", 2)
add_body(doc, "Pengembangan berikutnya disarankan berfokus pada sinkronisasi pipeline RDF, "
         "penambahan test suite dan CI, pemindahan varian lama ke area arsip, serta pembangunan "
         "algoritma rekomendasi dengan skor yang terukur. Dengan langkah tersebut, repository "
         "dapat berkembang dari demonstrasi teknologi menjadi sistem rekomendasi yang lebih "
         "andal, personal, dan mudah dipelihara.")

# Bibliography
doc.add_page_break()
add_heading(doc, "DAFTAR PUSTAKA", 1)
references = [
    "[1] World Wide Web Consortium. RDF 1.1 Concepts and Abstract Syntax. W3C Recommendation, 25 Februari 2014. https://www.w3.org/TR/rdf11-concepts/",
    "[2] World Wide Web Consortium. SPARQL 1.1 Query Language. W3C Recommendation, 21 Maret 2013. https://www.w3.org/TR/sparql11-query/",
    "[3] Apache Software Foundation. Apache Jena Fuseki Documentation. https://jena.apache.org/documentation/fuseki2/ (diakses 29 Juli 2026).",
    "[4] Streamlit. Streamlit Documentation. https://docs.streamlit.io/ (diakses 29 Juli 2026).",
    "[5] RDFLib Team. RDFLib Documentation. https://rdflib.readthedocs.io/en/stable/ (diakses 29 Juli 2026).",
    "[6] SPARQLWrapper Project. SPARQLWrapper Documentation. https://sparqlwrapper.readthedocs.io/en/stable/ (diakses 29 Juli 2026).",
    "[7] pandas Development Team. pandas.read_csv Documentation. https://pandas.pydata.org/pandas-docs/stable/reference/api/pandas.read_csv.html (diakses 29 Juli 2026).",
    "[8] The Movie Database. TMDB API Documentation. https://developer.themoviedb.org/docs/getting-started (diakses 29 Juli 2026).",
    "[9] Team Itihaad. Movie Recommendation System Using RDF and SPARQL, repository lokal dan dokumentasi proyek, working tree 29 Juli 2026.",
]
for ref in references:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.first_line_indent = Inches(-0.3)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    set_run_font(p.add_run(ref), size=10.5)

# Appendices
doc.add_page_break()
add_heading(doc, "LAMPIRAN A CONTOH KUERI SPARQL", 1)
add_heading(doc, "A.1 Menghitung Seluruh Triple", 2)
add_code(doc, """SELECT (COUNT(*) AS ?total)
WHERE {
  ?subject ?predicate ?object
}""")
add_heading(doc, "A.2 Mengambil Detail Film", 2)
add_code(doc, """PREFIX ex: <http://example.org/movies#>
SELECT ?title ?overview ?director ?actor_name
WHERE {
  ?movie ex:title "Spy Kids" ;
         ex:overview ?overview ;
         ex:directedBy ?director ;
         ex:hasActor ?actor .
  ?actor ex:title ?actor_name .
}""")
add_heading(doc, "A.3 Genre Terpopuler", 2)
add_code(doc, """PREFIX ex: <http://example.org/movies#>
SELECT ?genre (COUNT(?genre) AS ?count)
WHERE {
  ?movie ex:genre ?genre .
}
GROUP BY ?genre
ORDER BY DESC(?count)
LIMIT 10""")

doc.add_page_break()
add_heading(doc, "LAMPIRAN B KAMUS DATA FILM", 1)
add_table(doc, ["Kolom", "Makna", "Representasi RDF"], [
    ("movie_id", "ID film TMDB", "ex:movie_id dan bagian URI ex:Movie<ID>"),
    ("title", "Judul film", "ex:title / xsd:string"),
    ("overview", "Sinopsis ringkas", "ex:overview / xsd:string"),
    ("runtime", "Durasi dalam menit", "ex:runtime / xsd:float"),
    ("budget", "Anggaran produksi", "ex:budget / xsd:integer"),
    ("revenue", "Pendapatan film", "ex:revenue / xsd:integer"),
    ("original_language", "Bahasa asli", "ex:original_language / xsd:string"),
    ("release_date", "Tanggal rilis", "ex:release_date / xsd:date"),
    ("genre_1..5", "Hingga lima genre", "ex:genre / xsd:string"),
    ("Actor_1..3", "Tiga pemeran utama", "ex:hasActor / ex:Actor"),
    ("Director", "Sutradara", "ex:directedBy / ex:Director"),
    ("Country_1..5", "Negara produksi", "ex:producedIn / ex:Country"),
], [1900, 3150, 4310], font_size=8.7)

# Core properties
doc.core_properties.title = "Laporan Dokumentasi Movie Recommendation System Using RDF and SPARQL"
doc.core_properties.subject = "Analisis repository, arsitektur, implementasi, pengujian, dan panduan operasional"
doc.core_properties.author = "Team Itihaad"
doc.core_properties.keywords = "RDF, SPARQL, Apache Jena Fuseki, Streamlit, knowledge graph, movie recommendation"
doc.core_properties.comments = "Disusun berdasarkan working tree repository per 29 Juli 2026."

doc.save(OUTPUT)
print(OUTPUT)
